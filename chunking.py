"""
Loads and chunks documents from a knowledge folder, supporting .md, .pdf, and .docx.
To add a new document: just drop the file into the knowledge_docs/ folder and
push to GitHub — no code changes needed.

PDF pages are handled in three tiers, cheapest first:
  1. Pages with a real embedded text layer -> extracted directly (free, instant,
     perfectly accurate). Most "digital" PDF pages fall here.
  2. Pages with no text layer (i.e. the page is a scanned image) -> read with
     local Tesseract OCR (free, unlimited, runs on-device). Good for plain
     paragraphs, not reliable for complex multi-column diagrams.
  3. Pages you've explicitly flagged in VISION_PAGES below (e.g. flow-diagram
     pages that Tesseract garbles) -> transcribed with a free Groq vision
     model, which actually understands the diagram's layout. This is the
     only tier that costs API tokens, so it's reserved for the handful of
     pages that actually need it.
"""
import os
import re
import time
import base64
import numpy as np
from docx import Document
import pymupdf
import pytesseract
from PIL import Image
import io
from groq import Groq

KNOWLEDGE_DIR = "knowledge_docs"

# --- Tier 3 config: list the 1-indexed page numbers (per filename, as it
# appears in knowledge_docs/) that need the vision model — typically flow
# diagrams or other complex multi-column layouts that Tesseract can't read
# in the right order. Everything else is handled by tiers 1/2 for free.
# Example: "RMS_USER_MANUALv7.pdf": [44, 45, 67, 89]
VISION_PAGES = {
    # "RMS_USER_MANUALv7.pdf": [44],
}

MIN_NATIVE_TEXT_CHARS = 20  # below this, treat the page as having no real text layer

OCR_LANGUAGES = "eng"  # add "+chi_sim+chi_tra" if your PDFs contain Chinese text

VISION_MODEL = "qwen/qwen3.6-27b"  # Groq's free multimodal model, used only for VISION_PAGES

VISION_PROMPT = """Transcribe all text content from this page image into clean, well-structured plain text.

Rules:
- Preserve the reading order and logical structure exactly as a human reading the page would.
- If the page contains a flow diagram or process chart made of boxes/arrows arranged in a row or
  grid, describe it as a numbered sequence of stages in the correct order (left-to-right, or
  top-to-bottom if stacked). List ALL bullet points under each stage completely before moving on
  to the next stage — never interleave text from different boxes/columns.
- If the page contains a table, reproduce it as a markdown table.
- Preserve section headings and numbering exactly as shown (e.g. "4.4.1 EOP Flow").
- The page may contain a mix of English and other languages (e.g. Chinese) — transcribe each
  piece of text in its original language, do not translate.
- Do not add commentary, explanations, or any information not present in the image.
- Do not skip visible text, including footers, page numbers, or captions.
- Output only the transcribed content, nothing else — no preamble like "Here is the text:".
"""

# Reasoning-capable vision models can emit private work inside <think> tags.
# Treat an unclosed opening tag as extending to the end of the response: a
# truncated reasoning block is never valid manual content.
THINKING_BLOCK_RE = re.compile(r"<think\b[^>]*>.*?(?:</think\s*>|\Z)", re.IGNORECASE | re.DOTALL)
TRANSCRIPT_PAGE_MARKER_RE = re.compile(r"(<!--\s*Page\s+\d+\s*-->)", re.IGNORECASE)
def strip_thinking_content(text):
    """Remove model reasoning blocks before text can be chunked or embedded."""
    # A broken tag must not remove the remainder of a multi-page transcript.
    parts = TRANSCRIPT_PAGE_MARKER_RE.split(text or "")
    return "".join(
        part if TRANSCRIPT_PAGE_MARKER_RE.fullmatch(part) else THINKING_BLOCK_RE.sub("", part)
        for part in parts
    )


def _read_markdown(filepath):
    with open(filepath, "r", encoding="utf-8") as f:
        return strip_thinking_content(f.read())


def _ocr_page(page, min_gap_px=15, bin_px=5, line_tol=12):
    """
    Tier 2: local, free, unlimited OCR via Tesseract for image-only pages
    that are NOT in VISION_PAGES. Uses word-level bounding boxes to detect
    multi-column layouts and read them in the correct left-to-right,
    top-to-bottom order (plain image_to_string regularly scrambles this).

    This is a heuristic and not foolproof — it depends on there being a
    real empty gap (>= min_gap_px) between columns, so it can still fail on
    diagrams with narrow gutters or busy borders. If a page still comes out
    wrong after this, add its page number to VISION_PAGES instead of tuning
    these parameters further.

    Returns (text, column_count).
    """
    pix = page.get_pixmap(dpi=250)
    img_bytes = pix.tobytes("png")
    img = Image.open(io.BytesIO(img_bytes))

    data = pytesseract.image_to_data(img, lang=OCR_LANGUAGES, output_type=pytesseract.Output.DICT)
    words = []
    for i in range(len(data["text"])):
        text = data["text"][i].strip()
        if not text:
            continue
        left, top, w = data["left"][i], data["top"][i], data["width"][i]
        words.append({"text": text, "left": left, "top": top, "right": left + w})

    if not words:
        return "", 0

    min_left = min(w["left"] for w in words)
    max_right = max(w["right"] for w in words)
    n_bins = (max_right - min_left) // bin_px + 1
    density = np.zeros(n_bins, dtype=int)
    for w in words:
        b0, b1 = (w["left"] - min_left) // bin_px, (w["right"] - min_left) // bin_px
        density[b0:b1 + 1] += 1

    empty = density == 0
    boundaries = []
    run_start = None
    for i, is_empty in enumerate(empty):
        if is_empty and run_start is None:
            run_start = i
        elif not is_empty and run_start is not None:
            if (i - run_start) * bin_px >= min_gap_px:
                boundaries.append((run_start + i) // 2 * bin_px + min_left)
            run_start = None
    boundaries = [min_left - 1] + boundaries + [max_right + 1]

    columns = [[] for _ in range(len(boundaries) - 1)]
    for w in words:
        for ci in range(len(boundaries) - 1):
            if boundaries[ci] <= w["left"] < boundaries[ci + 1]:
                columns[ci].append(w)
                break
    columns = [col for col in columns if col]  # drop the empty gaps themselves

    page_lines = []
    for col in columns:
        col_sorted = sorted(col, key=lambda w: w["top"])
        line_groups = []
        for w in col_sorted:
            placed = False
            for g in line_groups:
                if abs(g[0]["top"] - w["top"]) < line_tol:
                    g.append(w)
                    placed = True
                    break
            if not placed:
                line_groups.append([w])
        for g in line_groups:
            g_sorted = sorted(g, key=lambda w: w["left"])
            page_lines.append(" ".join(w["text"] for w in g_sorted))
        page_lines.append("")

    # column_count is returned alongside the text so callers can use it as an
    # automatic signal for "this page likely needs vision-tier transcription"
    # (a real multi-column gap is exactly the layout Tesseract garbles) —
    # see _transcribe_pdf_auto, which escalates on this instead of requiring
    # a human to hand-curate VISION_PAGES.
    return "\n".join(page_lines).strip(), len(columns)


def _vision_ocr_page(page, client, dpi=200, max_retries=3):
    """
    Tier 3: renders a page to an image and asks a vision-capable LLM (via
    Groq's free API) to transcribe it, preserving reading order and
    structure. Only used for pages listed in VISION_PAGES, since each call
    costs a meaningful chunk of Groq's free daily token budget for this
    model — running it on every page of a large PDF will exhaust that
    budget well before the document is done (each page costs roughly
    3,500-4,700 tokens; the free tier is 200,000 tokens/day for this model).
    """
    pix = page.get_pixmap(dpi=dpi)
    png_bytes = pix.tobytes("png")
    b64_image = base64.b64encode(png_bytes).decode("utf-8")

    for attempt in range(max_retries):
        try:
            response = client.chat.completions.create(
                model=VISION_MODEL,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": VISION_PROMPT},
                            {
                                "type": "image_url",
                                "image_url": {"url": f"data:image/png;base64,{b64_image}"},
                            },
                        ],
                    }
                ],
                temperature=0,
                reasoning_effort="none",
                reasoning_format="hidden",
                max_completion_tokens=4096,
            )
            return strip_thinking_content(response.choices[0].message.content).strip()
        except Exception as e:
            if attempt == max_retries - 1:
                print(f"  Warning: vision transcription failed for a page after {max_retries} attempts: {e}")
                return ""
            wait_s = 10 * (attempt + 1)
            print(f"  Vision API call failed ({e}), retrying in {wait_s}s...")
            time.sleep(wait_s)


def _read_pdf(filepath, groq_api_key=None):
    """
    Extracts text from a PDF page by page, using the cheapest tier that
    works for each page (see module docstring). A Groq API key is only
    required if this PDF has pages listed in VISION_PAGES.
    """
    filename = os.path.basename(filepath)
    vision_page_nums = set(VISION_PAGES.get(filename, []))

    doc = pymupdf.open(filepath)
    pages_text = []
    client = None  # created lazily, only if a vision page is actually hit

    for i, page in enumerate(doc):
        page_num = i + 1

        native_text = page.get_text().strip()
        if len(native_text) >= MIN_NATIVE_TEXT_CHARS:
            pages_text.append(native_text)
            continue

        if page_num in vision_page_nums:
            if client is None:
                if not groq_api_key:
                    groq_api_key = os.environ.get("GROQ_API_KEY")
                if not groq_api_key:
                    raise ValueError(
                        f"{filename} has page {page_num} listed in VISION_PAGES, which "
                        "requires a Groq API key. Pass groq_api_key=... to "
                        "load_and_chunk_knowledge_folder(), or set the GROQ_API_KEY "
                        "environment variable."
                    )
                client = Groq(api_key=groq_api_key)
            print(f"  {filename} page {page_num}: no text layer, using vision model (configured)...")
            text = _vision_ocr_page(page, client)
        else:
            print(f"  {filename} page {page_num}: no text layer, using local OCR...")
            text, _ = _ocr_page(page)

        pages_text.append(text)

    doc.close()
    return strip_thinking_content("\n\n".join(pages_text))


# Hard ceilings for the fully-automatic upload pipeline (see
# _transcribe_pdf_auto). Unlike _read_pdf above — which relies on a human
# manually curating VISION_PAGES ahead of time — this path decides for
# itself, per page, whether vision transcription is warranted, with no
# human reviewing the outcome before it becomes live chatbot content. These
# caps exist because that combination (auto-decide + no review + public,
# ungated upload) has no other backstop against runaway Groq usage or
# multi-minute uploads: a single huge or adversarial PDF must not be able
# to exhaust the shared daily vision-token budget or hang the request
# indefinitely by itself.
AUTO_UPLOAD_MAX_TOTAL_PAGES = 150  # keeps worst-case (all-OCR) processing time bounded
AUTO_UPLOAD_MAX_VISION_PAGES = 30  # ~30 * 4,500 tokens =~ 135K of the 200K/day vision budget

# A page with a diagram/table pasted in as a picture can still have PLENTY
# of native text (headings, footers, step labels sitting around the image)
# — enough to sail past MIN_NATIVE_TEXT_CHARS while the actual diagram
# content itself is invisible to page.get_text() entirely, since it lives
# only inside the image. On this project's own manual, the known diagram
# pages (e.g. "4.4.1 EOP Flow", "6.6.1 Virement Flow") have 300-800+ chars
# of surrounding native text — well past that threshold — so gating
# escalation on native-text length alone would silently never trigger on
# exactly the pages this feature exists for. A large embedded image is
# checked FIRST and independently, regardless of how much native text is
# also on the page.
SIGNIFICANT_IMAGE_AREA_FRACTION = 0.15


def _page_has_significant_image(page, area_fraction=SIGNIFICANT_IMAGE_AREA_FRACTION):
    """True if an embedded image covers at least `area_fraction` of the
    page — a small logo/icon stays well under this (tested at ~6% on this
    project's own manual), while a pasted-in diagram/table/screenshot is
    typically 30%+."""
    page_area = page.rect.width * page.rect.height
    if not page_area:
        return False
    total = 0.0
    for img in page.get_images(full=True):
        try:
            for rect in page.get_image_rects(img[0]):
                total += rect.width * rect.height
        except Exception:
            continue  # a malformed/inline image reference shouldn't crash the page
        if total / page_area >= area_fraction:
            return True
    return False


def _transcribe_pdf_auto(
    pdf_bytes,
    groq_api_key,
    max_total_pages=AUTO_UPLOAD_MAX_TOTAL_PAGES,
    max_vision_pages=AUTO_UPLOAD_MAX_VISION_PAGES,
    progress_callback=None,
):
    """
    Fully-automatic version of the PDF tier system for uploads that will
    never get a human curating VISION_PAGES or spot-checking the output
    afterward — the tier-3 (vision) escalation decision is made BY THE CODE
    per page instead of by a human, using two independent signals:

      - A large embedded image (see _page_has_significant_image) — checked
        first, regardless of native text length, since a diagram pasted in
        as a picture can coexist with plenty of unrelated native text.
      - Failing that, a page with no native text layer that Tesseract's own
        column-gap detection finds genuinely multi-column (see _ocr_page) —
        catches scanned (not text-layer) diagram pages the same way.

    Either signal escalates the page to vision, up to max_vision_pages for
    this document. A plain single-column scanned paragraph page (the case
    tier 2 already handles fine) is deliberately NOT escalated — vision
    calls are reserved for pages that actually need them, so the budget
    isn't spent before it reaches genuinely hard pages later in the file.

    progress_callback(page_num, total_pages, tier_used: 1|2|3), if given, is
    invoked after each page — used by the caller to render a progress bar,
    since a large PDF can take several minutes end to end.

    Returns (markdown_text, stats) where stats is a dict with
    total_pages / vision_pages_used / ocr_pages, so the caller can tell the
    user what actually happened.
    """
    doc = pymupdf.open(stream=pdf_bytes, filetype="pdf")
    total_pages = len(doc)
    if total_pages > max_total_pages:
        doc.close()
        raise ValueError(
            f"PDF has {total_pages} pages, over the {max_total_pages}-page limit for "
            "automatic in-app processing. Split it into smaller files, or transcribe it "
            "locally with transcribe_full_pdf.py and upload the resulting .md instead."
        )

    client = None  # created lazily, only if a page is actually escalated to vision
    vision_pages_used = 0
    ocr_pages = 0
    pages_text = []

    for i, page in enumerate(doc):
        page_num = i + 1
        native_text = page.get_text().strip()
        wants_vision = _page_has_significant_image(page)

        if not wants_vision and len(native_text) >= MIN_NATIVE_TEXT_CHARS:
            pages_text.append(native_text)
            tier_used = 1
        else:
            if not wants_vision:
                # No big image, but also no usable native text (e.g. a
                # scanned page) — check whether it's genuinely multi-column
                # before deciding this needs vision too.
                ocr_text, column_count = _ocr_page(page)
                wants_vision = column_count >= 2
            else:
                ocr_text = None  # only computed lazily below if vision isn't available/fails

            if wants_vision and vision_pages_used < max_vision_pages and groq_api_key:
                if client is None:
                    client = Groq(api_key=groq_api_key)
                vision_text = _vision_ocr_page(page, client)
            else:
                vision_text = None

            if vision_text:
                pages_text.append(vision_text)
                vision_pages_used += 1
                tier_used = 3
            else:
                # No vision (unavailable, over budget, or the call failed
                # after its own retries) — fall back to the best text we
                # can still get: local OCR, or native text if that's all
                # there is.
                if ocr_text is None:
                    ocr_text, _ = _ocr_page(page)
                pages_text.append(ocr_text or native_text)
                ocr_pages += 1
                tier_used = 2

        if progress_callback:
            progress_callback(page_num, total_pages, tier_used)

    doc.close()

    parts = [f"<!-- Page {n} -->\n\n{strip_thinking_content(text)}" for n, text in enumerate(pages_text, start=1)]
    markdown_text = "\n\n---\n\n".join(parts)
    stats = {
        "total_pages": total_pages,
        "vision_pages_used": vision_pages_used,
        "ocr_pages": ocr_pages,
        "native_text_pages": total_pages - vision_pages_used - ocr_pages,
    }
    return markdown_text, stats


def _read_docx(filepath):
    doc = Document(filepath)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            style = (para.style.name or "").lower()
            if "heading 1" in style:
                parts.append(f"# {para.text}")
            elif "heading 2" in style:
                parts.append(f"## {para.text}")
            elif "heading 3" in style:
                parts.append(f"### {para.text}")
            else:
                parts.append(para.text)
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        if rows:
            parts.append("\n".join(rows))
    return "\n\n".join(parts)


def _read_txt(filepath):
    with open(filepath, "r", encoding="utf-8", errors="replace") as f:
        return strip_thinking_content(f.read())


LOADERS = {
    ".md": _read_markdown,
    ".txt": _read_txt,
    ".pdf": _read_pdf,
    ".docx": _read_docx,
}


PAGE_MARKER_RE = re.compile(r"<!--\s*Page\s+\d+\s*-->", re.IGNORECASE)
HEADER_LINE_RE = re.compile(r"^#{1,4}\s+(.+)$", re.MULTILINE)
# Vision-transcribed pages use inconsistent header styling from page to page
# (some use "**Section 3.1.1:**" bold, others use a plain bulleted/quoted
# line like '*   "7.5 Goods Receipt and Asset Tagging"' with no bold at
# all) — since the model's formatting choice isn't reliable, the most
# robust signal is the section NUMBER pattern itself ("7.5 Some Title"),
# which the document's actual structure guarantees is consistent.
#
# The title must start with a real WORD ([A-Z][a-z], e.g. "Application") —
# not just any capital letter. A bare capital letter followed by non-letters
# is exactly the shape of a unit abbreviation on a spec-table row ("1.8 L
# (approx. 12 cups)", "3.3 V input") — those parse as "section 1.8, title
# starting with L" under the old pattern and get torn out of the table they
# belong to mid-row. Confirmed on a real uploaded spec table: this is what
# silently broke "Water tank capacity" (label) from its own value.
SECTION_LINE_RE = re.compile(r"^\d+\.\d+(?:\.\d+)?\s+[A-Z][a-z].{1,69}$")
# Generated manuals do not consistently use Markdown headings, but their
# numbered RMS section headings are reliable boundaries. Split these before
# the size fallback so Purchasing cannot share an embedding with the
# neighbouring Pay-and-Claim (<RM1K) section.
SECTION_SPLIT_RE = re.compile(
    r"(?=^(?:#{1,4}\s+)?\d+\.\d+(?:\.\d+)?\s+[A-Z][a-z][^\n]{1,119}$)",
    re.MULTILINE,
)
BOLD_HEADER_RE = re.compile(r"^\*\*([^*]{3,80})\*\*", re.MULTILINE)

# FAQ files (e.g. "Frequently Asked Questions - RMS.txt") are a flat numbered
# list of "N. <question ending in ?>" followed by an (unlabeled) answer —
# confirmed against the actual files in knowledge_docs/, neither of which
# uses literal "Q:"/"A:" labels at all (an earlier version of this pattern
# required them, so it never matched either real file — the whole FAQ file
# silently fell through to the generic splitter/character-slicer below
# instead, undetected because nothing errored). Without this, dozens of
# unrelated Q&A pairs get embedded together as a single vector (or worse,
# blindly sliced mid-answer) — the retriever can find "the FAQ file" but not
# "this specific question", and the model then has to guess which of the
# many answers buried in that blob is the right one. Splitting on the "N.
# <question>?" boundary gives each Q&A pair its own chunk, titled with the
# question text itself, so retrieval can match the actual question being asked.
FAQ_QA_SPLIT_RE = re.compile(r"(?=^\s*\d+\.\s+.+\?\s*$)", re.MULTILINE)
# Non-greedy up to the first "?" so this stops at the end of the question
# even if the answer that follows contains further numbered/punctuated text.
FAQ_Q_LINE_RE = re.compile(r"^\s*\d+\.\s+(.+?\?)", re.MULTILINE | re.DOTALL)


def _looks_like_faq(text):
    """At least 3 numbered '<question>?' entries is a reasonable bar for
    'this file is a FAQ list', low enough to catch short files but high
    enough to not misfire on a manual that happens to number a couple of
    unrelated questions."""
    return len(FAQ_QA_SPLIT_RE.findall(text)) >= 3


def _split_faq_into_chunks(text, source_name):
    blocks = FAQ_QA_SPLIT_RE.split(text)
    blocks = [b.strip() for b in blocks if b.strip()]
    result = []
    for block in blocks:
        m = FAQ_Q_LINE_RE.search(block)
        title = re.sub(r"\s+", " ", m.group(1)).strip() if m else _extract_chunk_title(block, source_name)
        result.append({"title": title, "text": block, "source": source_name})
    return result


def _find_section_heading_line(chunk_text):
    """
    Looks for a line that IS ENTIRELY a numbered section heading (e.g.
    "7.5 Goods Receipt and Asset Tagging"), after stripping common bullet/
    quote wrapping — not just a section number mentioned in passing inside
    a longer sentence (e.g. "...see 7.2.1 Pre-Approval claim)." would NOT
    match, since that's a cross-reference, not this chunk's own heading).
    """
    for line in chunk_text.split("\n"):
        stripped = line.strip().lstrip("*-").strip()
        stripped = stripped.strip('"').strip()
        if SECTION_LINE_RE.match(stripped):
            return stripped
    return None


# Repeated letterhead lines that appear at the top of most pages in this
# manual's vision transcription — never a real title, so the title-fallback
# below skips past them looking for the first line that actually says
# something distinctive.
_TITLE_BOILERPLATE_LINES = {"mmu", "multimedia university", "rms manual"}


def _extract_chunk_title(chunk_text, source_name):
    m = HEADER_LINE_RE.search(chunk_text)
    if m:
        return m.group(1).strip()
    heading = _find_section_heading_line(chunk_text)
    if heading:
        return heading
    # Prefer a bold pseudo-header that looks like a real section marker
    # (contains a digit, e.g. "Section 3.1.1") over generic model preamble
    # like "**Analyze the image:**" or "**List:**" that sometimes appears
    # first in a vision-transcribed page.
    for m in BOLD_HEADER_RE.finditer(chunk_text):
        candidate = m.group(1).strip().rstrip(":").strip()
        if re.search(r"\d", candidate):
            return candidate
    m = BOLD_HEADER_RE.search(chunk_text)
    if m:
        return m.group(1).strip().rstrip(":").strip()
    # Last resort: first non-empty line. Many pages open with the same
    # letterhead boilerplate ("MMU" / "MULTIMEDIA UNIVERSITY" / "RMS
    # Manual") before their actual content, which — with no better title
    # candidate above — made this fallback hand out the literal string
    # "MMU" as the title for 20 different, unrelated chunks. Useless in the
    # "Sources used" UI (tells the reader nothing) and hides which of
    # several same-titled chunks actually has the answer. Skip boilerplate
    # lines so this lands on the first line that actually says something.
    for line in chunk_text.split("\n"):
        line = line.strip()
        if not line or PAGE_MARKER_RE.match(line) or line.lower() in _TITLE_BOILERPLATE_LINES:
            continue
        return re.sub(r"^#+\s*", "", line)[:80]
    return source_name



# Minimum characters of BODY text a chunk needs beyond its own first line
# (usually its heading) to count as real, answerable content rather than a
# heading-only fragment (a Table of Contents entry, mainly — see the filter
# in _split_text_into_chunks). Deliberately small: real content in this
# project's documents is consistently much longer than this in practice, so
# this only needs to catch the "nothing at all follows the heading" case,
# not judge how much content is "enough".
MIN_CHUNK_BODY_CHARS = 10


def _split_text_into_chunks(text, source_name, max_chunk_chars=6000, overlap_chars=300):
    # Split on numbered section headings across the WHOLE document first, not
    # per-page. A table or flow-diagram enumeration that continues onto the
    # next page has no heading of its own on that second page, so splitting
    # by "<!-- Page N -->" markers before splitting by heading used to tear
    # it into two chunks — retrieval would then only surface half of it.
    all_sections = SECTION_SPLIT_RE.split(text)
    all_sections = [s.strip() for s in all_sections if s.strip()]
    if not all_sections:
        all_sections = [text]

    # A section with no numbered "N.N Title" heading of its own falls back to
    # "<!-- Page N -->" markers (inserted by transcribe_full_pdf.py /
    # _transcribe_pdf_auto) as a secondary boundary — regardless of size, not
    # only when oversized. This matters for any document that isn't numbered
    # in the RMS manual's "N.N"/"N.N.N" style at all (plenty of real uploads
    # use plain "1. Title" headings instead, which SECTION_SPLIT_RE can't
    # safely match — that pattern collides with ordinary numbered list items
    # like "1. Unbox the machine..." inside a procedure). Without this, such
    # a document has NO heading-level split points anywhere, so as long as
    # it's under max_chunk_chars it would end up as ONE chunk covering the
    # entire document — every unrelated section diluting the same embedding,
    # which measurably hurts retrieval for a specific fact (confirmed: a
    # short uploaded manual's spec-table answer stopped being found once it
    # was no longer split by page). A section that DOES have its own
    # recognized heading is left whole — that's what keeps a table or
    # flow-diagram spanning two pages from being torn apart.
    refined_sections = []
    for section in all_sections:
        has_own_heading = _find_section_heading_line(section) is not None
        if has_own_heading and len(section) <= max_chunk_chars:
            refined_sections.append(section)
            continue
        page_pieces = PAGE_MARKER_RE.split(section)
        page_pieces = [p.strip() for p in page_pieces if p.strip()]
        refined_sections.extend(page_pieces if len(page_pieces) > 1 else [section])
    all_sections = refined_sections

    # A whole section (a page, or a "#"-delimited region within a page) stays
    # as ONE chunk regardless of length, as long as it's under the safety
    # ceiling — this is what actually prevents mid-section cuts, rather than
    # hoping max_chunk_chars is "big enough". Character-slicing only kicks
    # in for the rare pathological section that exceeds even that ceiling,
    # and when it does, every slice keeps the section's own title attached
    # so a later slice never loses its identity (which is also embedded
    # together with the text in retriever.py — see there for why that matters).
    chunks = []  # list of (chunk_text, forced_title_or_None)
    for section in all_sections:
        if len(section) <= max_chunk_chars:
            chunks.append((section, None))
        else:
            title = _extract_chunk_title(section, source_name)
            start = 0
            first = True
            while start < len(section):
                end = start + max_chunk_chars
                piece = section[start:end]
                if not first:
                    piece = f"[continued — {title}]\n\n{piece}"
                chunks.append((piece, title))  # reuse the same clean title for every slice
                start = end - overlap_chars
                first = False

    result = []
    for chunk, forced_title in chunks:
        stripped = chunk.strip()
        if not stripped:
            continue
        # A chunk that's nothing but its own heading line — a Table of
        # Contents entry like "6.6.1 Virement Flow 77" (heading text + page
        # number, no body) is exactly this shape — provides zero answerable
        # content, yet still competes as a candidate in every retrieval
        # query. Worse, it can actively WIN: BM25 scores a heading-only
        # chunk higher than the real, complete chunk for that same
        # section, since almost all of its (very short) text IS the
        # matching terms — no body text dilutes that ratio. Confirmed
        # live: the TOC's "Virement Flow" line outranked the real 4-stage
        # chunk in production, so the model only ever saw the empty one.
        # 54 of this manual's 296 chunks (its whole Table of Contents) were
        # exactly this before this filter.
        first_line, _, rest = stripped.partition("\n")
        if len(rest.strip()) < MIN_CHUNK_BODY_CHARS:
            continue
        title = forced_title if forced_title else _extract_chunk_title(stripped, source_name)
        result.append({
            "title": title,
            "text": chunk,
            "source": source_name,
        })
    return result


def load_and_chunk_knowledge_folder(folder_path=KNOWLEDGE_DIR, groq_api_key=None):
    """Loads every supported file in the folder and returns combined chunks."""
    all_chunks = []
    if not os.path.isdir(folder_path):
        raise FileNotFoundError(
            f"Knowledge folder '{folder_path}' not found. "
            f"Create it and add your .md/.pdf/.docx files there."
        )

    files = sorted(os.listdir(folder_path))
    loaded_any = False
    for filename in files:
        ext = os.path.splitext(filename)[1].lower()
        if ext not in LOADERS:
            continue
        filepath = os.path.join(folder_path, filename)
        if ext == ".pdf":
            text = _read_pdf(filepath, groq_api_key=groq_api_key)
        else:
            text = LOADERS[ext](filepath)
        if not text.strip():
            continue
        if ext == ".txt" and _looks_like_faq(text):
            chunks = _split_faq_into_chunks(text, source_name=filename)
        else:
            chunks = _split_text_into_chunks(text, source_name=filename)
        all_chunks.extend(chunks)
        loaded_any = True

    if not loaded_any:
        raise FileNotFoundError(
            f"No supported files (.md, .pdf, .docx) found in '{folder_path}'."
        )
    return all_chunks


def load_and_chunk_markdown(filepath, max_chunk_chars=3500, overlap_chars=300):
    text = _read_markdown(filepath)
    return _split_text_into_chunks(text, source_name=os.path.basename(filepath),
                                    max_chunk_chars=max_chunk_chars, overlap_chars=overlap_chars)


if __name__ == "__main__":
    chunks = load_and_chunk_knowledge_folder()
    print(f"Total chunks: {len(chunks)}")
    sources = set(c["source"] for c in chunks)
    print(f"Loaded from files: {sources}")
